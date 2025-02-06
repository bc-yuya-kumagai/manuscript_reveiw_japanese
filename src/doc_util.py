# Wordファイルの解析に関するユーティリティ関数を提供するモジュール

from typing import List
from docx import Document
from docx.text.paragraph import Paragraph
from lxml import etree
from zipfile import ZipFile
from xml.etree import ElementTree as ET
from src.entity import Section
import re
import logging
import src.general_util as gu
import src.llm_util as llm
import src.kanji as kanji
logger = logging.getLogger(__name__)
# 問の見出しスタイルID
question_heading_style_id = 'af8'
# wordファイルから下線部のrunを抽出する
def get_underline_runs(doc,first_paragraph_index:int,last_paragraph_index:int):
    """ first_paragraph_indexとlast_paragraph_indexの間で underlined runを取得する
    last_paragraph_indexが-1の場合は最後のdocx_fileの最後の段落までを検索する"""
    runs = []
    for p in doc.paragraphs[first_paragraph_index:last_paragraph_index]:
        for run in p.runs:
            # スタイルが1-5-10または1-5-20の場合は下線部として抽出
            if (len(run.text)>0) and (run.style.style_id == '1-5-10' or run.style.style_id == '1-5-20'): #二重傍線スタイル1-5-20は誤爆する？
                runs.append(run)
    return runs

def check_countains_text(doc:Document,text:str, first_paragraph_index:int,last_paragraph_index:int):
    """ first_paragraph_indexとlast_paragraph_indexの間でtextが含まれているかを検索する
    last_paragraph_indexが-1の場合は最後のdocx_fileの最後の段落までを検索する"""
    for p in doc.paragraphs[first_paragraph_index:last_paragraph_index]:
        if text in p.text:
            return True
    return False

def get_previous_text_index_run(sideline_run):
  
    prv = sideline_run._element.getprevious()
    while prv !=  None and (prv.text == None or prv.text == '') :
        prv = prv.getprevious()

    if prv is not None:
        return prv
    else:
        return None

def get_first_question_paragraph_index(doc, start:int, end:int):
    """ 、question_heading_style_idの見出しを持つ段落のインデックスを返す
    インデックス順に文書が構成されていることを前提とする
    

    """
    # question_heading_style_idの見出しを持つ段落のインデックスを返す
    # インデックス順に文書が構成されていることを前提とする
    # for i, p in enumerate(doc.paragraphs):
    #         for r in p.runs:
    #             if r.style.style_id == question_heading_style_id:
    #                 return i

    # 仮の実装 インデントや空白なしで行頭が「問」で始まる段落を問の見出しとする
    for i, p in enumerate(doc.paragraphs[start:end+1]):
        if p.text.startswith("問"):
            return i
        
def get_questions(doc:Document,start:int, end:int)->List[Paragraph]:
    """docから問から次の問までのPhraseを取得する
    """
    questions=[]
    question_phrases = []
    for p in doc.paragraphs[start:end+1]:
        
        if p.text.startswith("問"):
            # 次の問に到達した場合は、現時点の設問文をリストに追加し、設問文を初期化する
            if len(question_phrases) > 0:
                questions.append(question_phrases)
            question_phrases=[p]

        elif len(question_phrases) > 0:
            question_phrases.append(p)
    questions.append(question_phrases)
    return questions

def extract_question_paragraphs(doc: Document, start:int, end:int) -> List[Paragraph]:
    """
    文書から「問」で始まるパラグラフを抽出する

    Returns:
        List[Paragraph]: 「問」で始まるParagraphオブジェクトのリスト
    """
    question_paragraphs = []
    for paragraph in doc.paragraphs[start:end+1]:
        # TODO: 問の取得条件を適切に設定する  
        if paragraph.text.startswith("問") and not paragraph.text.startswith("問題"):
            # 「問」で始まるパラグラフをリストに追加
            question_paragraphs.append(paragraph)
    return question_paragraphs


def clean_sileline_list_in_page_break(list):
    """リストから改ページで発生するごみを削除する。
    TODO: ちゃんと改ページを判断する

    """
    return [i for i in list if len(i.index_text)<4]


def get_paragraph_text_by_keyword(doc:Document,word:str):
    """docの段落からwordが含まれている段落のテキストを取得する
    """
    for p in doc.paragraphs:
        if word in p.text:
            yield p.text

def get_choice_indexes_from_choices_list(question_phrases:List[Paragraph]):
    """選択肢のリストから選択肢の添え字を取得する
    """
    indexes = []
    for p in question_phrases:
        # フレーズ内で <w:rStyle w:val="2-3-10"/>であるrunを取得する
        for run in p.runs:
            if run.style.style_id == '2-3-10':
                indexes.append(run.text)
        
    return indexes

def find_continuous_run_indices(paragraph:Paragraph, target:str):
    """
    Check if the target string exists as a sequence of characters across the elements
    of the string_list and return the indices where it is found.

    Args:
        string_list (list of str): List of strings to search.
        target (str): The target string to find.

    Returns:
        list of int: List of indices where the target string is found.
    """
    # Join the list into a single continuous string with a delimiter
  
    combined_text = "".join(r.text for r in paragraph.runs)
    if target not in combined_text or len(target) == 0:
        return []


    # Determine which indices of the original list are involved
    indices = []
    offset = 0
    for i, r in enumerate(paragraph.runs):
        for c in r.text:
            if len(c) == 0:
                continue
            if c == target[offset]:
               indices.append(i)
               offset += 1
               if offset == len(target): # offsetがtargetの長さに達したら、後続のtargetの検索を行うため、offsetを0に戻す
                     offset = 0
            else:
                offset = 0
    # indicesの要素をユニークにする
    return list(set(indices))


def get_explanation_of_questions(doc: Document) -> List[str]:
    """
    ドキュメントから設問の解説を抽出する。ただし、「解答・配点」に到達した時点で抽出を終了する。
    
    Args:
        doc (Document): docxファイルを読み込んだDocumentオブジェクト
    Returns:
        List[str]: 各設問の解説を含むリスト
    """
    explanation_flg = False  # 「●設問解説」を見つけたかどうか
    all_questions = []  # すべての設問を格納
    current_question = []  # 現在処理中の設問を格納

    for p in doc.paragraphs:
        text = p.text.strip()

        # 「●設問解説」が見つかったらフラグをオン
        if text.startswith("●設問解説"):
            explanation_flg = True
            continue

        # 「解答・配点」が見つかったらフラグをオフ
        if explanation_flg and "解答・配点" in text:
            explanation_flg = False
            continue

        # フラグがオンの場合、設問を処理
        if explanation_flg:
            # 新しい「問」で始まる設問が見つかったら保存
            if text.startswith("問") and current_question:
                all_questions.append("\n".join(current_question))
                current_question = []

            # 現在の段落を追加
            if text:
                current_question.append(text)

    # 最後の設問を保存
    if current_question:
        all_questions.append("\n".join(current_question))

    return all_questions
def font_analyzer(docx_file_path: str, paragraph: Paragraph):
    """
    段落内のテキストとフォント情報を解析する関数。

    Args:
        docx_file_path (str): 対象のWord文書のファイルパス。
        paragraph (Paragraph): 段落オブジェクト（`python-docx` の Paragraph クラス）。

    Returns:
        List[Dict[str, Any]]: 段落内のテキストとフォント情報のリスト。
        各辞書は以下の形式を持つ:
            - "text" (str): 実際のテキスト。
            - "font" (str): フォント名。
    """
    buffer = []  # 段落内のテキスト情報を一時的に保持する
    for run in paragraph.runs:
        theme_font = None
        r_element = run._element
        # <w:rPr> 要素を検索
        rPr = r_element.find(".//w:rPr", namespaces=r_element.nsmap)
        if rPr is not None:
            # テーマ情報 (<w:rFonts> の属性) を取得
            rFonts = rPr.find(".//w:rFonts", namespaces=r_element.nsmap)
            if rFonts is not None:
                theme_font = rFonts.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}asciiTheme")
                if theme_font and theme_font.startswith("minor"):
                    theme_font = find_theme_font_schemas(docx_file_path)["minorFont"]["Jpan"]
                elif theme_font and theme_font.startswith("major"):
                    theme_font = find_theme_font_schemas(docx_file_path)["majorFont"]["Jpan"]

        # フォントヒエラルキー
        font_name = None
        if run.font.name is None and theme_font is None:
            font_name = get_style_by_id(docx_file_path, run.style.style_id)["font"]["ascii"]
        elif theme_font:
            font_name = theme_font
        else:
            font_name = run.font.name

        buffer.append({
            "text": run.text,
            "font": font_name,
        })
    return buffer

def get_style_by_id(docx_file_path: str, style_id: str) -> dict:
    """
    styles.xml 内の指定された styleId の設定をオブジェクト形式で返す関数。

    :param docx_file_path: WORDファイルのパス
    :param style_id: 取得したいスタイルの ID (styleId)
    :return: スタイル設定を格納した辞書オブジェクト
    """
    # styles.xml を取得
    with ZipFile(docx_file_path) as docx:
        styles_xml = docx.read("word/styles.xml")
    # 名前空間の定義
    NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    # styles.xml の読み込み
    tree = etree.XML(styles_xml)
  

    # 指定された styleId を持つ <w:style> を検索
    style_element = tree.find(f".//w:style[@w:styleId='{style_id}']", namespaces=NS)
    if style_element is None:
        return {"error": f"Style ID '{style_id}' not found."}

    # スタイル情報を辞書に格納
    style_info = {
        "styleId": style_id,
        "type": style_element.attrib.get(f"{{{NS['w']}}}type", "Not Specified"),
        "customStyle": style_element.attrib.get(f"{{{NS['w']}}}customStyle", "0"),
        "name": None,
        "basedOn": None,
        "link": None,
        "font": {"ascii": None, "hAnsi": None, "eastAsia": None, "hint": None},
        "color": None,
        "size": None,
    }

    # スタイル名
    name_element = style_element.find(f"{{{NS['w']}}}name", namespaces=NS)
    if name_element is not None:
        style_info["name"] = name_element.attrib.get(f"{{{NS['w']}}}val")

    # 継承元スタイル
    based_on_element = style_element.find(f"{{{NS['w']}}}basedOn", namespaces=NS)
    if based_on_element is not None:
        style_info["basedOn"] = based_on_element.attrib.get(f"{{{NS['w']}}}val")

    # リンクされたスタイル
    link_element = style_element.find(f"{{{NS['w']}}}link", namespaces=NS)
    if link_element is not None:
        style_info["link"] = link_element.attrib.get(f"{{{NS['w']}}}val")

    # ランプロパティ <w:rPr> を探索
    rpr_element = style_element.find(f"{{{NS['w']}}}rPr", namespaces=NS)
    if rpr_element is not None:
        # フォント情報
        rfonts_element = rpr_element.find(f"{{{NS['w']}}}rFonts", namespaces=NS)
        if rfonts_element is not None:
            style_info["font"]["ascii"] = rfonts_element.attrib.get(f"{{{NS['w']}}}ascii")
            style_info["font"]["hAnsi"] = rfonts_element.attrib.get(f"{{{NS['w']}}}hAnsi")
            style_info["font"]["eastAsia"] = rfonts_element.attrib.get(f"{{{NS['w']}}}eastAsia")
            style_info["font"]["hint"] = rfonts_element.attrib.get(f"{{{NS['w']}}}hint")

        # 色情報
        color_element = rpr_element.find(f"{{{NS['w']}}}color", namespaces=NS)
        if color_element is not None:
            style_info["color"] = color_element.attrib.get(f"{{{NS['w']}}}val")

        # フォントサイズ
        size_element = rpr_element.find(f"{{{NS['w']}}}sz", namespaces=NS)
        if size_element is not None:
            style_info["size"] = size_element.attrib.get(f"{{{NS['w']}}}val")

    return style_info


def find_theme_font_schemas(word_file_path):
    """
    Wordファイルから主要フォントスキームと副次フォントスキームを抽出します。

    Args:
        word_file_path (str): Word (.docx) ファイルのパス。

    Returns:
        dict: 'majorFont' と 'minorFont' のキーを持つ辞書。各キーには、それぞれのフォントマッピングが含まれます。
    """
    # WordファイルはZIP形式で圧縮されているため、ZipFileを使用します
    with ZipFile(word_file_path, 'r') as docx:
        # テーマ情報は"word/theme/theme1.xml"に格納されています
        theme_path = 'word/theme/theme1.xml'
        if theme_path not in docx.namelist():
            raise FileNotFoundError(f"{theme_path} not found in the Word file.")

        # テーマXMLファイルを読み取ります
        theme_xml = docx.read(theme_path)

    # XMLを解析します
    root = ET.fromstring(theme_xml)
    namespaces = {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}

    # majorFontとminorFontを探します
    font_scheme = root.find('.//a:fontScheme', namespaces)
    if font_scheme is None:
        raise ValueError("Font scheme not found in the theme XML.")

    major_font = {}
    minor_font = {}

    # majorFontの抽出
    major_font_element = font_scheme.find('a:majorFont', namespaces)
    if major_font_element is not None:
        for font in major_font_element:
            script = font.attrib.get('script', 'latin')
            typeface = font.attrib.get('typeface', '')
            major_font[script] = typeface

    # minorFontの抽出
    minor_font_element = font_scheme.find('a:minorFont', namespaces)
    if minor_font_element is not None:
        for font in minor_font_element:
            script = font.attrib.get('script', 'latin')
            typeface = font.attrib.get('typeface', '')
            minor_font[script] = typeface

    return {
        'majorFont': major_font,
        'minorFont': minor_font
    }

def is_start_section(paragraph: Paragraph) -> bool:
    """
    大問の開始を判定する関数。

    Args:
        paragraph (Paragraph): python-docx の Paragraph オブジェクト。

    Returns:
        bool: 大問の開始であれば True、そうでなければ False。
    """
    if len(paragraph.text.strip()) == 0:
        return False
    # 大問の開始をスタイルで判定
    daimon_runs = [r for r in paragraph.runs if (r.text.strip() != "" ) and (r.style is not None) and (r.style.style_id == "1-20")]
    return len(daimon_runs) > 0

def is_end_section(paragraph: Paragraph) -> bool:
    """
    大問の終了を判定する関数。
    大問の終わりには特別なマーカーは存在しないので常にFalse を返す。

    Args:
        paragraph (Paragraph): python-docx の Paragraph オブジェクト。

    Returns:
        bool: False 固定。
    """
    return paragraph.text.strip() == "END"
def extract_sections(doc: Document) -> List[Section]:
    """
    問題の文書から大問を抽出する関数。

    Args:
        doc (Document): python-docx の Document オブジェクト。

    Returns:
        List[Section]: 抽出されたセクションのリスト。
    """
    sections:List[Section] = []
    section_intervals = gu.extract_intervals(doc.paragraphs, is_start=is_start_section, is_end=is_end_section)
    for interval in section_intervals:
        section:Section = Section(section_number=interval.items[0].text[0], body_text="\n".join(p.text for p in interval.items[1:]))
        section.star_paragraph_index = interval.start
        section.end_paragraph_index = interval.end
        sections.append(section)
    
    return sections

def is_kanjinumchar(char:str):
    """漢数字かどうかを判定する"""
    kanji_num_chars = ["〇","一","二","三","四","五","六","七","八","九","十","百","千"]
    return char in kanji_num_chars

def is_start_expl_section(paragraph: Paragraph) -> bool:
    """
    解説の大問の開始を判定する関数。

    Args:
        paragraph (Paragraph): python-docx の Paragraph オブジェクト。

    Returns:
        bool: 解説の大問の開始であれば True、そうでなければ False。
    """
    # 漢数字の後に、"現代文"、"古文"、"漢文"が続く場合に大問の開始と判定
    if len(paragraph.text)> 20:
        False
    title_offset = paragraph.text.find("現代文")
    if len(title_offset) == -1:
        title_offset = paragraph.text.find("古文")
    if len(title_offset) == -1:
        title_offset = paragraph.text.find("漢文")
    
    if title_offset == -1:
        return False
    
    if is_kanjinumchar(paragraph.text[:title_offset].strip()[-1]):
        return True

def is_end_explsection(paragraph: Paragraph) -> bool:
    return False

def extract_explain_sections(doc: Document) -> List[Section]:
    """
    解説の文書から大問を抽出する関数。

    Args:
        doc (Document): python-docx の Document オブジェクト。

    Returns:
        List[Section]: 抽出されたセクションのリスト。
    """
    sections:List[Section] = []
    section_intervals = gu.extract_intervals(doc.paragraphs, is_start=is_start_expl_section, is_end=is_end_explsection)
    for interval in section_intervals:
        section:Section = Section(section_number=interval.items[0].text[0], body_text="\n".join(p.text for p in interval.items[1:]))
        section.star_paragraph_index = interval.start
        section.end_paragraph_index = interval.end
        sections.append(section)
    
    return sections

def kanji_number_to_arabic_number(kanji_numr:str):
    """漢数字をアラビア数字に変換する"""
    # 漢数字をアラビア数字に変換
    kanji_to_number = {
        "〇": 0,
        "一": 1,
        "二": 2,
        "三": 3,
        "四": 4,
        "五": 5,
        "六": 6,
        "七": 7,
        "八": 8,
        "九": 9,
    }
    arabic_number = "".join(str(kanji_to_number[char]) for char in kanji_numr if char in kanji_to_number)
    return arabic_number

title_question = re.compile(r'^(?:【[^】]+】)?[一二三四五六七八九十百千]+　[^\s　]+(?:　[^\s　]+)*')
def extract_question_number(doc):
    """設問番号を抽出する"""
    question_main_score_list = []
    found = False
    question_text = ""

    for p in doc.paragraphs:
        text = p.text.strip()
        # 設問タイトルにマッチしたら、新しい設問の開始
        if title_question.match(text):
            # すでに収集した設問があれば処理（次の設問の開始前に保存）
            if question_text:
                question_main_score_list.append(llm.extract_main_score_from_text(question_text))

            # 新しい設問の収集を開始
            found = True
            question_text = text  # 設問のタイトル部分をセット
            continue

        # 空行（改行のみ）の場合、設問収集を終了
        if text == "":
            found = False
            continue

        # 設問の収集中なら続けて追加
        if found:
            question_text += "\n" + text

    # 最後の設問を処理
    if question_text:
        question_main_score_list.append(llm.extract_main_score_from_text(question_text))

    return question_main_score_list
title_question = re.compile(r'^[一二三四五六七八九十百千]+　[^\s　]+(?:　[^\s　]+)*')
def extract_main_text(doc: Document, start, end) -> list[list[Document]]:
    """大問から、問の前までの全テキストを抽出します。"""
    all_texts = []
    current_text = []
    start_collecting = False

    for p in doc.paragraphs[start:end+1]:
        if title_question.match(p.text):
            if current_text:  # 既に収集中の本文があるなら保存
                all_texts.append(current_text)
                current_text = []  # 次の本文のためにリセット
            start_collecting = True  # 新しい本文の収集開始

        if start_collecting:
            if p.text.startswith("問"):
                if current_text:  # 最後の収集データを保存
                    all_texts.append(current_text)
                start_collecting = False  # 収集終了
                current_text = []
                continue
            
            current_text.append(p)

    # 最後の本文があれば追加
    if current_text:
        all_texts.append(current_text)

    return all_texts

annotation_pattern = re.compile(r"^[０-９]+　+[^\s　]+(?:[-―－…・]+[^\s　]+)*(?:　+|[-―－…・])+.+")
def extract_annotation_text_to_list(annotation_paragraph: Document) -> list[str]:
    """本文のパラグラフから傍注のリストを作成して返す関数"""
    annotation_names = []
    is_collecting_annotations = False

    for line in annotation_paragraph:
        if line.text.startswith("（注）"):  # 「（注）」が始まったら収集を開始
            is_collecting_annotations = True  # 傍注の収集開始
        if is_collecting_annotations:
            for line in line.text.split("\n"):
                
                line = line.lstrip()
                if line.startswith("（注）"):
                    line = line[3:]  # 「（注）」を削除
                parts = line.split("――")  # 全角スペースで分割
                if len(parts) > 1:  # 2つ以上の要素があるか確認
                    annotation_names.append(parts[0].split("　")[1])  # 2つ目の要素を取得
                else:
                    is_collecting_annotations = False  # フォーマット違いの行が出たら収集停止
                    break  # ループを抜ける

    return annotation_names

def extract_main_text_and_annotation_to_main_text(documents_list: Document) -> list[Document]:
    """大問から、問の前までの全テキストを抽出しますから問題本文を抜き出します。"""
    main_text_list = []
    for p in documents_list:
        if p.text.startswith("（注）"):
            break
        main_text_list.append(p)
            
    return main_text_list
def get_explanation_of_questions(doc: Document,start:int, end:int) -> List[str]:
    """
    解説ドキュメントから設問の解説を抽出する。ただし、「解答・配点」に到達した時点で抽出を終了する。
    
    Args:
        doc (Document): docxファイルを読み込んだDocumentオブジェクト
    Returns:
        List[str]: 各設問の解説を含むリスト
    """
    explanation_flg = False  # 「●設問解説」を見つけたかどうか
    all_questions = []  # すべての設問を格納
    current_question = []  # 現在処理中の設問を格納

    for p in doc.paragraphs[start:end+1]:
        text = p.text.strip()

        # 「●設問解説」が見つかったらフラグをオン
        if text.startswith("●設問解説"):
            explanation_flg = True
            continue

        # 「解答・配点」が見つかったらフラグをオフ
        if explanation_flg and "解答・配点" in text:
            explanation_flg = False
            continue

        # フラグがオンの場合、設問を処理
        if explanation_flg:
            # 新しい「問」で始まる設問が見つかったら保存
            if text.startswith("問") and current_question:
                all_questions.append("\n".join(current_question))
                current_question = []

            # 現在の段落を追加
            if text:
                current_question.append(text)

    # 最後の設問を保存
    if current_question:
        all_questions.append("\n".join(current_question))

    return all_questions
def kanji_number_to_arabic_number(kanji_numr:str):
    """漢数字をアラビア数字に変換する"""
    # 漢数字をアラビア数字に変換
    kanji_to_number = {
        "〇": 0,
        "一": 1,
        "二": 2,
        "三": 3,
        "四": 4,
        "五": 5,
        "六": 6,
        "七": 7,
        "八": 8,
        "九": 9,
    }
    arabic_number = "".join(str(kanji_to_number[char]) for char in kanji_numr if char in kanji_to_number)
    return arabic_number

title_question = re.compile(r'^(?:【[^】]+】)?[一二三四五六七八九十百千]+　[^\s　]+(?:　[^\s　]+)*')
def extract_question_number(doc):
    """設問番号を抽出する"""
    question_main_score_list = []
    found = False
    question_text = ""

    for p in doc.paragraphs:
        text = p.text.strip()
        # 設問タイトルにマッチしたら、新しい設問の開始
        if title_question.match(text):
            # すでに収集した設問があれば処理（次の設問の開始前に保存）
            if question_text:
                question_main_score_list.append(llm.extract_main_score_from_text(question_text))

            # 新しい設問の収集を開始
            found = True
            question_text = text  # 設問のタイトル部分をセット
            continue

        # 空行（改行のみ）の場合、設問収集を終了
        if text == "":
            found = False
            continue

        # 設問の収集中なら続けて追加
        if found:
            question_text += "\n" + text

    # 最後の設問を処理
    if question_text:
        question_main_score_list.append(llm.extract_main_score_from_text(question_text))

    return question_main_score_list
title_question = re.compile(r'^[一二三四五六七八九十百千]+　[^\s　]+(?:　[^\s　]+)*')
def extract_main_text(doc: Document, start:int, end:int) -> List[Paragraph]:
    """大問から、問の前までの全テキストを抽出します。"""
    all_texts = []
    current_text = []
    start_collecting = False

    for p in doc.paragraphs[start:end+1]:
        if title_question.match(p.text):
            if current_text:  # 既に収集中の本文があるなら保存
                all_texts.append(current_text)
                current_text = []  # 次の本文のためにリセット
            start_collecting = True  # 新しい本文の収集開始

        if start_collecting:
            if p.text.startswith("問"):
                if current_text:  # 最後の収集データを保存
                    all_texts.append(current_text)
                start_collecting = False  # 収集終了
                current_text = []
                continue
            
            current_text.append(p)

    # 最後の本文があれば追加
    if current_text:
        all_texts.append(current_text)

    return all_texts

annotation_pattern = re.compile(r"^[０-９]+　+[^\s　]+(?:[-―－…・]+[^\s　]+)*(?:　+|[-―－…・])+.+")
def extract_annotation_text_to_list(annotation_paragraph: Document) -> list[str]:
    """本文のパラグラフから傍注のリストを作成して返す関数"""
    annotation_names = []
    is_collecting_annotations = False

    for line in annotation_paragraph:
        if line.text.startswith("（注）"):  # 「（注）」が始まったら収集を開始
            is_collecting_annotations = True  # 傍注の収集開始
        if is_collecting_annotations:
            for line in line.text.split("\n"):
                
                line = line.lstrip()
                if line.startswith("（注）"):
                    line = line[3:]  # 「（注）」を削除
                parts = line.split("――")  # 全角スペースで分割
                if len(parts) > 1:  # 2つ以上の要素があるか確認
                    annotation_names.append(parts[0].split("　")[1])  # 2つ目の要素を取得
                else:
                    is_collecting_annotations = False  # フォーマット違いの行が出たら収集停止
                    break  # ループを抜ける

    return annotation_names

def extract_main_text_and_annotation_to_main_text(documents_list: Document) -> list[Document]:
    """大問から、問の前までの全テキストを抽出しますから問題本文を抜き出します。"""
    main_text_list = []
    for p in documents_list:
        if p.text.startswith("（注）"):
            break
        main_text_list.append(p)
            
    return main_text_list
def get_explanation_of_questions(doc: Document, start:int, end:int) -> List[str]:
    """
    解説ドキュメントから設問の解説を抽出する。ただし、「解答・配点」に到達した時点で抽出を終了する。
    
    Args:
        doc (Document): docxファイルを読み込んだDocumentオブジェクト
    Returns:
        List[str]: 各設問の解説を含むリスト
    """
    explanation_flg = False  # 「●設問解説」を見つけたかどうか
    all_questions = []  # すべての設問を格納
    current_question = []  # 現在処理中の設問を格納

    for p in doc.paragraphs[start:end+1]:
        text = p.text.strip()

        # 「●設問解説」が見つかったらフラグをオン
        if text.startswith("●設問解説"):
            explanation_flg = True
            continue

        # 「解答・配点」が見つかったらフラグをオフ
        if explanation_flg and "解答・配点" in text:
            explanation_flg = False
            continue

        # フラグがオンの場合、設問を処理
        if explanation_flg:
            # 新しい「問」で始まる設問が見つかったら保存
            if text.startswith("問") and current_question:
                all_questions.append("\n".join(current_question))
                current_question = []

            # 現在の段落を追加
            if text:
                current_question.append(text)

    # 最後の設問を保存
    if current_question:
        all_questions.append("\n".join(current_question))

    return all_questions

# 
from src.check import InvalidItem

def set_section_at_invalid_iterms(invalid_items:List[InvalidItem], section_number:str):
    """invalid_itemのsection_numberを修正する"""
    results = []
    for i in invalid_items:
        i.section_number = section_number
        results.append(i)
    return results

def has_ruby(run):
    # Run要素のXMLを取得
    xml = run._r.xml
    # ルビ要素の存在を確認
    tree = etree.fromstring(xml)
    namespace = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    ruby_elements = tree.findall(f'.//{namespace}ruby')
    return len(ruby_elements) > 0

def get_runs_with_not_ordinary_kanji_without_ruby(paragraphs:List[Paragraph]):
    """段落から通常の漢字以外の文字を含むrunを取得する"""
    runs_with_not_ordinary_kanji = []
    for paragraph in paragraphs:
        for run in paragraph.runs:
            for char in run.text:
                # charの文字コードが漢字の帯域ある場合は、そのrunをリストに追加
                if 0x4E00 <= ord(char) <= 0x9FFF and char not in kanji.ORDINARY_KANJI_SET:
                    if not has_ruby(run) :
                        runs_with_not_ordinary_kanji.append((char, run))
                        break

    return runs_with_not_ordinary_kanji

# 使用例
if __name__ == "__main__":
    word_file_path = "example.docx"  # Wordファイルのパスを指定
    font_scheme = find_theme_font_schemas(word_file_path)
    print("Major Font:", font_scheme['majorFont'])
    print("Minor Font:", font_scheme['minorFont'])
