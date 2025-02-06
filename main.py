from typing import List, Generator
from fastapi import FastAPI, UploadFile, File, HTTPException, status
from fastapi.responses import HTMLResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
import logging
import os
from src import doc_util
from src import check as ck
from src.check import InvalidItem, SideLine
from src.entity import Section

app = FastAPI()

# CORSの設定（必要に応じて）
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ロギングの設定を最初に行う
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# temp_problem_file_path, temp_solution_file_path
def analyze_problem_doc(problem_doc, temp_problem_file_path):
    # 問題のみのチェック
    problem_invalid_list = []
    sections:List[Section] = doc_util.extract_sections(problem_doc)
    for section in sections:
        # 問の見出しが最初に始まる箇所を特定 <- これが、問題文と設問の境界になる
        first_question_paragraph_index:int = doc_util.get_first_question_paragraph_index(problem_doc,start = section.star_paragraph_index, end=section.end_paragraph_index)
        if not first_question_paragraph_index:
            return {"errors":[{"type":"INDEX_NOT_FOUND","message":"問の見出しスタイルIDが見つかりませんでした"}]}
        
        # 傍線部取得（問題文の中から傍線部のrunを取得）
        passage_side_line_runs = doc_util.get_underline_runs(problem_doc, 0, first_question_paragraph_index-1)

        passage_sideLine_list = []
        for run in passage_side_line_runs:
            passage_sideLine_list.append( SideLine(index_text=doc_util.get_previous_text_index_run(run).text, passage=run.text))

        # ページ区切り箇所にゴミが残るのでそれを削除
        passage_sideLine_list = doc_util.clean_sileline_list_in_page_break(passage_sideLine_list)
        # 傍線部の添え字重複チェック
        problem_invalid_list += doc_util.set_section_at_invalid_iterms(ck.check_duplicated_index(passage_sideLine_list), section.section_number)
    
        # 問のテキストを設問ごとにリストでの取得
        question_texts = doc_util.get_questions(problem_doc, start= section.star_paragraph_index, end=section.end_paragraph_index)

        # 非常用漢字にルビが振られていることのチェック
        check_not_ordinary_kanji_without_ruby_results = ck.check_not_ordinary_kanji_without_ruby(problem_doc, start=section.star_paragraph_index, end=section.end_paragraph_index)
        for error in check_not_ordinary_kanji_without_ruby_results:
            error.section_number = section.section_number
            problem_invalid_list.append(error)

        # 選択肢設問の設問文で、「適切」ではなく「適当」となっているかチェックし、適切ならエラーを返す
        check_keyword_exact_match_in_question_results = ck.check_keyword_exact_match_in_question(question_texts)
        errors = doc_util.set_section_at_invalid_iterms(check_keyword_exact_match_in_question_results, section_number=section.section_number)
        problem_invalid_list.extend(list(errors))

        # 傍線部の連番飛びチェック
        jumped = ck.check_jumped_index(passage_sideLine_list)
        if isinstance(jumped,InvalidItem):
            jumped.section_number = section.section_number
            problem_invalid_list.append(jumped)

        # 傍線部の添え字が設問内で参照されているかチェック
        slideline_questions = list(doc_util.get_paragraph_text_by_keyword(problem_doc, "傍線部"))
        result_sl_mapping = ck.check_mapping_sileline_index_userd_in_questions(passage_sideLine_list, slideline_questions)
        if isinstance(result_sl_mapping, InvalidItem):
            result_sl_mapping.section_number = section.section_number
            problem_invalid_list.append(result_sl_mapping)

        # 設問内の添字が問題文中にあるかチェック
        result_sl_mapping = ck.check_mapping_sileline_index_appear_in_passage(passage_sideLine_list, slideline_questions)
        if isinstance(result_sl_mapping, InvalidItem):
            result_sl_mapping.section_number = section.section_number
            problem_invalid_list.append(result_sl_mapping)

        # 選択肢のチェック
        for q_idx, question in enumerate(question_texts):
            question_text = "\n".join([q.text for q in question])
            if ck.get_question_type(question_text) == "選択式":
                for i in doc_util.set_section_at_invalid_iterms(ck.check_choices_mapping(question), section_number=section.section_number):
                    i.question_number = q_idx + 1
                    problem_invalid_list.append(i)
        
        # 選択肢に重複や歯抜けがないかチェック
        for q_idx, question in enumerate(question_texts):
            question_text = "\n".join([q.text for q in question])
            if ck.get_question_type(question_text) == "選択式":
                error = ck.check_choices_sequence(question)
                if isinstance(error, InvalidItem):
                    error.section_number = section.section_number
                    error.question_number = q_idx + 1
                    problem_invalid_list.append(error)
                
        # 「適当でないもの」がMSゴシックであるかチェック
        for q_idx, question in enumerate(question_texts):
            result_check_font_of_unfit_item = ck.check_font_of_unfit_item(question)
            if isinstance(result_check_font_of_unfit_item, InvalidItem):
                result_check_font_of_unfit_item.section_number = section.section_number
                result_check_font_of_unfit_item.question_number = q_idx + 1
                problem_invalid_list.append(result_check_font_of_unfit_item)

        # 選択肢設問の設問文で、「適切」ではなく「適当」となっているかチェックし、適切ならエラーを返す
        check_keyword_exact_match_in_question_statement = ck.check_keyword_exact_match_in_question(question_texts)
        if isinstance(check_keyword_exact_match_in_question_statement, InvalidItem):
            check_keyword_exact_match_in_question_statement.section_number = section.section_number
            problem_invalid_list.append(check_keyword_exact_match_in_question_statement)
    
        # 文書から「問」で始まるパラグラフを抽出する(各問の書き出しを取得する)
        extract_paragraphs = doc_util.extract_question_paragraphs(problem_doc, start=section.star_paragraph_index, end=section.end_paragraph_index)

        # 「問~」がMSゴシックかチェック
        check_heading_question_font_item = ck.check_heading_question_font(temp_problem_file_path, extract_paragraphs)
        if isinstance(check_heading_question_font_item, InvalidItem):
            check_heading_question_font_item.section_number = section.section_number
            problem_invalid_list.append(check_heading_question_font_item)

        # 設問番号が順番通りになっているかチェック
        check_kanji_number_orders =  doc_util.set_section_at_invalid_iterms(ck.check_kanji_question_index_order(extract_paragraphs), section_number=section.section_number)
        for error in check_kanji_number_orders:
            problem_invalid_list.append(error)

        #傍注の説明の内容が本文に入っているかチェック
        check_exists_annotation_result = ck.check_exists_annotation(problem_doc, start=section.star_paragraph_index, end=section.end_paragraph_index)
        if isinstance(check_exists_annotation_result, InvalidItem):
            check_exists_annotation_result.section_number = section.section_number
            problem_invalid_list.append(check_exists_annotation_result)

        # 設問の漢字書き取り問題に指定されたフレーズが含まれているかチェック
        check_writing_kanji_phrase_errors:Generator[InvalidItem] = ck.check_phrase_in_kanji_writing_question(question_texts)
        # check_writing_kanji_phrase_errorsの各要素にsection_numberを設定してproblem_invalid_listに追加
        for error in check_writing_kanji_phrase_errors:
            error.section_number = section.section_number
            problem_invalid_list.append(error)

        # 漢字読み取り問題時に、「（現代仮名遣いでよい。）」というフレーズが使われているかチェック
        check_kanji_reading_missing_result = ck.check_kanji_reading_missing_expressions(question_texts)
        if isinstance(check_kanji_reading_missing_result, InvalidItem):
            check_kanji_reading_missing_result.section_number = section.section_number
            problem_invalid_list.append(check_kanji_reading_missing_result)
            


    return problem_invalid_list

def analyze_solution_doc(solution_doc):
    explain_invalid_list = []
    sections = doc_util.extract_explain_sections(solution_doc)
    for section in sections:
        # 解説のみのチェック

        # 解説中に正答番号を指すものに対して、正答というフレーズが正しく使用されているか確認する。
        check_explanation_of_questions_error = ck.check_explanation_of_questions_include_word(solution_doc, start=section.star_paragraph_index, end=section.end_paragraph_index)
        if isinstance(check_explanation_of_questions_error, InvalidItem):
            explain_invalid_list.append(check_explanation_of_questions_error)
        
        # 記述設問の際、解説のポイントが存在しているかチェック
        check_answer_point = ck.check_answer_contains_points(solution_doc, start=section.star_paragraph_index, end=section.end_paragraph_index)
        if isinstance(check_answer_point, InvalidItem):
            check_answer_point.
            explain_invalid_list.append(check_answer_point)

        # ●設問解説ブロッック内の現代語訳部分の表記が、現代語訳ブロックに存在するかチェック
#        check_modern_translation = ck.check_modern_translation(solution_doc, start=section.star_paragraph_index, end=section.end_paragraph_index)
    return explain_invalid_list

def analyze_common_doc(problem_doc, solution_doc):
    common_invalid_list = []
    # 問題と解説両方をチェック
    if problem_doc and solution_doc:
        
        problem_texts = doc_util.get_questions(problem_doc, start=0, end=len(problem_doc.paragraphs))
        solution_texts = doc_util.get_questions(solution_doc, start=0, end=len(solution_doc.paragraphs))

        # 大問の配点をチェックする。
        part_question_score_check = ck.check_part_question_score(problem_doc, solution_doc)
        if isinstance(part_question_score_check, InvalidItem):
            common_invalid_list.append(part_question_score_check)

        # 問題文で文字数について言及されているものと解説文の文字数が一致しているかチェック
        check_question_and_answer_word_count=ck.check_question_sentence_word_count(problem_texts, solution_texts)
        if isinstance(check_question_and_answer_word_count, InvalidItem):
            common_invalid_list.append(check_question_and_answer_word_count)
    return common_invalid_list

def analyze_docx(temp_problem_file_path, temp_solution_file_path):
    """
    docx_file_pathで指定されたWordファイルを分析し、
    不備リストとメッセージを返す関数。
    """
    problem_doc = []
    solution_doc = []
    if temp_problem_file_path:
        problem_doc = Document(temp_problem_file_path)
    if temp_solution_file_path:
        solution_doc = Document(temp_solution_file_path)

    invalid_list = dict()

    if problem_doc:
        invalid_list["problem"] = analyze_problem_doc(problem_doc,temp_problem_file_path=temp_problem_file_path)

    if solution_doc:
        invalid_list["solution"] = analyze_solution_doc(solution_doc)

    if problem_doc and solution_doc:
        invalid_list["common"] = analyze_common_doc(problem_doc, solution_doc)

    return invalid_list


@app.get("/", response_class=HTMLResponse)
async def home_page():
    # 簡易的なアップロードフォーム
    return """
        <!DOCTYPE html>
        <html lang="ja">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1">
            <title>国語原稿チェックツール</title>
            <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
            <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/js/bootstrap.bundle.min.js"></script>
            <style>
                .upload-box {
                    border: 2px dashed #0d6efd;
                    border-radius: 10px;
                    padding: 100px;
                    text-align: center;
                    cursor: pointer;
                    transition: all 0.3s;
                }
                .upload-box:hover {
                    background-color: #f8f9fa;
                }
            </style>
            <script>
                function updateFileName(inputId, labelId) {
                    const input = document.getElementById(inputId);
                    const label = document.getElementById(labelId);
                    if (input.files.length > 0) {
                        label.textContent = input.files[0].name;
                    } else {
                        label.textContent = '.docx ファイルのみ対応';
                    }
                }

                function validateForm(event) {
                    const problemFile = document.getElementById('problem_file').files.length;
                    const solutionFile = document.getElementById('solution_file').files.length;
                    
                    if (problemFile === 0 && solutionFile === 0) {
                        alert('問題ファイルまたは解説ファイルのどちらかを選択してください。');
                        event.preventDefault();
                    }
                }
            </script>
        </head>
        <body class="bg-light">
            <div class="container text-center mt-5">
                <h1 class="mb-4">国語原稿チェックツール</h1>
                <form action="/upload" method="post" enctype="multipart/form-data" onsubmit="validateForm(event)">
                    <div class="row justify-content-center">
                        <div class="col-md-5">
                            <label class="upload-box d-block" for="problem_file">
                                <img src="https://cdn-icons-png.flaticon.com/512/1086/1086534.png" width="50" class="mb-2">
                                <p class="mb-1">問題ファイルをアップロード</p>
                                <small class="text-muted" id="problem_label">.docx ファイルのみ対応</small>
                            </label>
                            <input id="problem_file" name="problem_file" type="file" accept=".docx" class="d-none" onchange="updateFileName('problem_file', 'problem_label')">
                        </div>
                        <div class="col-md-5">
                            <label class="upload-box d-block" for="solution_file">
                                <img src="https://cdn-icons-png.flaticon.com/512/1086/1086534.png" width="50" class="mb-2">
                                <p class="mb-1">解説ファイルをアップロード</p>
                                <small class="text-muted" id="solution_label">.docx ファイルのみ対応</small>
                            </label>
                            <input id="solution_file" name="solution_file" type="file" accept=".docx" class="d-none" onchange="updateFileName('solution_file', 'solution_label')">
                        </div>
                    </div>
                    <button type="submit" class="btn btn-primary mt-4">確認する</button>
                </form>
            </div>
        </body>
        </html>
    """

async def save_temp_file(docx_file: UploadFile) -> str:
    """一時ファイルとしてdocxファイルを保存し、ファイルパスを返す"""
    temp_file_path = f"temp_{docx_file.filename}"
    with open(temp_file_path, "wb") as f:
        f.write(await docx_file.read())
    return temp_file_path

def delete_temp_file(file_path: str):
    """一時ファイルを削除する"""
    if file_path is None:
        return
    try:
        os.remove(file_path)
    except OSError as e:
        logger.error(f"Error deleting file {file_path}: {e}")

def convert_to_html_table(data):

    html = """
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/css/bootstrap.min.css" rel="stylesheet">
    
    """
    if data.get("problem"):

        if len(data["problem"]) == 0:
            html += """
                
                    <p >問題チェックOK! </p>
                
            """
        else:
            html += """問題のチェック結果
            <table class="table table-bordered table-striped table-hover">
        <thead class="table-dark">
            <tr>
                <th>大問</th>
                <th>問</th>
                <th>エラー種別</th>
                <th>メッセージ</th>
            </tr>
        </thead>
        <tbody>
        """
        for problem in data["problem"]:
            html += f"""
                <tr>
                    <td>{problem.section_number}</td>
                    <td>{problem.question_number}</td>
                    <td>{problem.type}</td>
                    <td>{problem.message}</td>
                </tr>
               
            """
        html += """
            </tbody>
        </table>"""


    if data.get("solution"):
        if len(data["solution"]) == 0:
            html += """
               <p>解説チェックOK!</p>
            """
        else:
            html += """解説のチェック結果
             <table class="table table-bordered table-striped table-hover">
        <thead class="table-dark">
            <tr>
                <th>大問</th>
                <th>問</th>
                <th>エラー種別</th>
                <th>メッセージ</th>
            </tr>
        </thead>
        <tbody>
            """
        for solution in data["solution"]:
            html += f"""
                <tr>
                    <td>{solution.section_number}</td>
                    <td></td>
                    <td>{solution.type}</td>
                    <td>{solution.message}</td>
                </tr>
            """
        html += """
            </tbody>
        </table>"""
        
    if data.get("common"):
        if len(data["common"]) == 0:
            html += """
            <p>問題/解説比較チェックOK! </p>
              
            """
        else:
            html += """問題/解説比較チェック結果
             <table class="table table-bordered table-striped table-hover">
        <thead class="table-dark">
            <tr>
                <th>大問</th>
                <th>問</th>
                <th>エラー種別</th>
                <th>メッセージ</th>
            </tr>
        </thead>
        <tbody>"""
        for common in data["common"]:
            html += f"""
                <tr>
                    <td></td>
                    <td></td>
                    <td>{common.type}</td>
                    <td>{common.message}</td>
                </tr>
            """
        
        html += """
            </tbody>
        </table>
    <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/js/bootstrap.bundle.min.js"></script>
    """
    return html


@app.post("/upload",response_class=HTMLResponse)
async def check_docx(problem_file: UploadFile = File(None), solution_file: UploadFile = File(None)):
    temp_problem_file_path = None
    temp_solution_file_path = None
    if problem_file and len(problem_file.filename) > 0:
        temp_problem_file_path = await save_temp_file(problem_file)
    if solution_file and len(solution_file.filename) > 0:
        temp_solution_file_path = await save_temp_file(solution_file)

    try:
        # 分析実行
        result = analyze_docx(temp_problem_file_path, temp_solution_file_path)
    finally:
        # 一時ファイル削除
        if problem_file:
            delete_temp_file(temp_problem_file_path)
        if solution_file:
            delete_temp_file(temp_solution_file_path)

    html_table = convert_to_html_table(result)
    return HTMLResponse(html_table)
