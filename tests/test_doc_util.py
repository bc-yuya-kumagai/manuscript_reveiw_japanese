from src.doc_util import get_style_by_id, kanji_number_to_arabic_number, extract_question_number, extract_annotation_text_to_list, extract_main_text_and_annotation_to_main_text
from unittest.mock import MagicMock
from unittest.mock import MagicMock
from docx.text.paragraph import Paragraph
from src.doc_util import get_style_by_id
import unittest
from docx import Document
from src.doc_util import extract_sections
from src.entity import Section


# テスト対象のサンプルファイルパス
SAMPLE_DOCX_PATH = "tests/resources/スタイル付_【問題A】自動原稿整理PoC_サンプル原稿（指摘箇所コメント付）.docx"





class TestDocUtil(unittest.TestCase):

    def setUp(self):
        # テスト用のドキュメントを作成
        self.doc = Document()
        self.doc.add_paragraph("一 大門1")
        self.doc.add_paragraph("これは大門1の内容です。")
        self.doc.add_paragraph("二 大門2")
        self.doc.add_paragraph("これは大門2の内容です。")

    def test_extract_sections(self):
        sections = extract_sections(self.doc)
        self.assertEqual(len(sections), 2)
        self.assertEqual(sections[0].section_number, "一")
        self.assertEqual(sections[0].body_text, "これは大門1の内容です。")
        self.assertEqual(sections[1].section_number, "二")
        self.assertEqual(sections[1].body_text, "これは大門2の内容です。")
    
    def test_extract_sections_from_docx(self):
        doc = Document(SAMPLE_DOCX_PATH)
        sections = extract_sections(doc)
        self.assertEqual(len(sections), 1)
        self.assertEqual(sections[1].section_number, "二")
        self.assertEqual(sections[1].body_text, "次の【文章Ⅰ】は「夢の子供」（一九九九年発表）の一節である。地質研究所に勤める「私」（カナ）は、大学院生である本間泉と暮らしている。「私」の母親の麻子は二ヵ月ほど前に急逝した。現在、「私」は風邪をこじらせて微熱が引かないまま勤めを休んでいる。以下はそれに続く場面である。これを読んで、後の問いに答えよ。（配点\u3000六〇）\n\n【文章Ⅰ】\n私が身動きをしたので、泉はふり返ったが唇に指を一本当ててみせ、「〈翠たち〉が来ているよ」と小声で教えた。\n「ほんとに？」と半身を起こすと、ベランダの手すりにレタスの葉をもぎとって並べたように、二羽のインコがとまっている。雨の滴が羽毛にはじかれて、全身にビーズのように附着している。今日はガラス戸が開いていたので、いくらか警戒気味の様子だ。いつもならⓐしごく無雑作にヒマワリの種に突進するのに、小首をかしげてためらっている。しかし私たちが静かにしていたので、やがてえさ台に飛びうつった。\nインコたちがこのマンションの周辺に出没しはじめたのは春先からだった。麻子が身体の異常を訴えて入院したころである。見舞いに行ったとき、自分の母親を励ますことに慣れていない私は話題に窮し、途中で見かけたこの鳥の話をした。\n「マンションの庭の木にね、大きなインコが二羽とまっていたのよ。羽の色はエメラルド・グリーンで、首のまわりに赤い輪があって、くちばしも真っ赤。エスニック調ですごくきれいなわりに、鳴き声はヤスリをこすり合わせるみたいなの」と言って私はそのいやな声を真似してみせた。ギイ、ギイ、ギーッと。\nそれから鳥の目は気味悪い、といつか彼女が言っていたことを思いだし、はっとして黙った。しかし麻子はもの憂げに私のほうに顔を向けて微笑んだうえ、意外なことを言った。\n「らしてみたら？\u3000ベランダにえさを置いてみたら？\u3000小さいころ、カナはよく地面にお米をまいて、スズメを集めてたでしょう」\n「そうだったね。じゃ、やってみようか」\n「うん、やってみて。そして写真とってきてちょうだい。私も見たいから。そんな妖艶な姿の鳥」\nそれよりか早く元気になって実物を見にくれば、と言いたかったが、(ｱ)ついに口から出すことができなかった。死の影に半分以上侵されている麻子にとって、も本気もたいして価値はちがわないのに、慰めだけが暗い病室のかな明りかもしれないのに。私はときどき自分の性格が嫌いになる。\nでもその日、病院の帰りに私は彼女の意に添おうとし、駅ビルのペットショップに寄ってインコの好物をたずねた。その店で円盤に支柱のついたえさ台と五キログラム入りのヒマワリの種の袋を求めると、タクシーに乗ってマンションに運びこんだ。大学から戻った泉が目を丸くするころには、ベランダにはインコを招きよせる装置がそろっていたのだ。\nマンションの庭に出入りしていた二羽のインコがヒマワリの種に気づくまでに一週間、ベランダに入って好物をついばむ勇気を獲得するまでにさらに一週間かかった。私たちは昼間ほとんど家にいなかったから、それはえさ台上の種子のピラミッドが崩れていくことで知れたのである。ベランダのコンクリートに散らばる白黒の甲虫に似た殻のふえ方で、インコたちがしだいに大胆になっていく様子がわかった。そしてとうとうある土曜日の朝、異様な叫びに驚いた私がこっそりカーテンのすきまからのぞいてみると、〈翠たち〉が円盤にのっていた。目覚めたばかりのあやふやな視野に収まった鮮やかな羽色のインコは、一瞬（注１）ゴーガンの風景画から飛びたってきたようであった。二羽のインコはくつろいでいるふうで、交互に首を下げ、長い足指とに曲がったくちばしを使って器用に殻をむいて食べた。私は泉にも見せたくなって揺り起こそうとしたが、あまりにぐっすりと眠っているので放っておいた。（注２）一週間分の夜勤の疲れを解消するのは、土曜日しかないのであったから。\n離れているときに呼びあう声の調子やおたがいの何気ない仕草から、私は二羽が単なる同類以上の関係にあると信じることができた。どちらが雌か雄か、見かけではわからなかったので、私たちはまとめて〈翠たち〉と名づけてしまった。\nやがてインコたちは室内に人がいようといまいと、平気でベランダに現われるようになった。もちろん無人のときも、空腹でありさえすれば飛んできていただろう。仕事帰りの私が、マンションのエレベーターの前で通いの管理人に呼びとめられたのはそのころのことである。\n「あのねえ、本間さん（彼は私が「潮野」という自分の姓を泉のそれと並べて郵便箱に貼りつけてあるのに、知らないふりをするのだ）、お宅の下の階から苦情が出ているんですけどねえ」\n「はい、どういうことで？」\nマンションに住むということは、(ｲ)苦情の渦の縁に住むことである。住民のだれかが渦に吸いこまれる前に調整するのが管理人の役目でもあることは、私も承知している。\n「上からヒマワリの種の殻が降ってくるそうですよ。ときには鳥のふんも。心当たりあります？」\n私は大急ぎであやまった。\n「はい、わかりました。以後注意します」\n「そうしてください」話が速やかについて彼はほっとしたらしい。「実は私も小鳥が好きでしてねえ」と話しはじめた。「家でジュウシマツを飼っているんですがね、いやあれは実によく卵を産みますなあ。ひなが育つと初めは近所に差しあげてたんですが、今では妻が里親探しに大変なんですよ。でも小鳥はかわいい。心がなごみます」\n「でもあのインコはべつに飼ってるわけではないんです」と私は誤解されると困るので、あわてて事情を説明した。「どこかの飼鳥が野生化したらしいんです。向うが勝手に遊びにきてる、ということなんです」\nわかっている、というふうに管理人はうなずいた。そしてさらに私の知らないことを教えてくれた。\n「あの鳥はもともとこの町内の保育園から集団脱走した数羽の子孫らしいんです。現在は二百羽ぐらいにふえて、工業大学の大イチョウを夜のねぐらにしているんですよ。東南アジアの原産らしいけれど、順応性が強いんでしょう。近所の家から通報があったとかで、このあいだＴＶニュースでもねぐら入りを放映してましたよ」という口調はどこか誇らしげであった。\n「そうだったんですか。きっとさぞ騒がしいでしょうねえ」\n私は〈翠たち〉の鳴き声の百倍の音響効果を想像して、つぶやいた。\n「あれがカナリアみたいな声だったらよかったのに。天は二物を与えなかったんですねえ。ふっふっふ」と管理人は含み声で笑った。\nエレベーターの中で私はめずらしく(ｲ)親しげな態度を示した管理人の話をした。彼の語った〝籠ぬけ鳥〟伝説を信用すると、この町には一組の〈翠たち〉どころか、すでに何十組もの〈翠たち〉が住んでいるのだ。そういえばカラスともハトともちがう鳥の群れが、雑居ビルの屋上をかすめて工業大学の方角に飛んでいく光景を見たことがある。夕闇に紛れて羽色まではわからなかったが、その折りふっとなじみのないものに会ったような違和感を受けた。あれはきっとインコの群れであったのだろう。\n部屋に戻ってからも、まだ先ほどの会話が頭について離れなかった。うちの〈翠たち〉の体内に、たぶん熱帯の森で捕獲された数代前の先祖の遺伝子があるのだと知って、何となく興奮したのだ。近くで聞けば不協和音の一種としか思えないあの鳴き声も、この町とは比べものにならないほど広大な生息地では仲間同士の通信に役だつことだろう。強い陽光のもとでエメラルド色の羽毛はより輝き、首のもようは火の輪のように燃えるだろう。この町内で生まれ育った〈翠たち〉が、過度に目だつ外観以上の何かを体内に受け継いでいるとしても不思議はない。そう思うとあの黒と赤の奇妙な二重眼が私たちの部屋に何を見ているのか、知りたい気がする。\nインコたちは私や泉が室内から見ていても平気でえさを食べつづけるほどれてきたが、カメラを向けるとあわてて逃げだした。光る目玉のようなものに対して、本能的な恐怖心を抱いているようだ。インコの写真を見たいとせがんだ麻子は、それが実現する前にに陥ってしまった。えづけは彼女のアイディアだったのに、私はとうにそのことを忘れ、彼女がいなくなってからも〈翠たち〉に夢中になっている。\nインコたちが満足して飛び去ると、私はえさをもらうひな鳥のように首をのばしてたずねた。「今日は何？」\n「三色ずしと春雨サラダとクラブサンドイッチ」と泉はあくびをしながら言った。「それからガラス戸閉めていいかな。雨足が強くなったから」\n食物と器をテーブルに並べるのは彼にまかせて、私は自分で戸を閉めにいった。インコの去ったあとのベランダには殻が散乱しているが、管理人の注意を受けてから手すりの下方に細い金網を張りめぐらせたので下の階には落ちていかないはずだ。餌台のほかには、アロエやハーブの鉢が何種類も並んでいる。泉がこまめに世話をするので、葉の色つやもよく茎は幹みたいに力強く立っている。マンションのどの階のどの部屋のベランダよりも鬱蒼と茂っている。インコたちには草原（注１）ッののシュのように映るかもしれない。\n生活を共にする前までは、泉に庭師の一面があることを知らなかった。熱にうかされていた一週間、(ｳ)私はベランダを眺めるたびにうねる緑の海原を連想していた。\n\n（注）１\u3000ゴーギャン\u3000――ポール・ゴーギャン（一八四八―一九〇三）。フランスの画家。\n\u3000\u3000\u3000２\u3000一週間分の夜勤\u3000――泉は深夜から明け方までコンビニエンスストアでアルバイトをしていて、アルバイトが終わるといつもそのお店の食料品を持って家に帰っ\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000てくる。\n\u3000\u3000\u3000３\u3000ブッシュ\u3000――低木の茂み。やぶ。\n\n問一\u3000二重傍線部ⓐ・ⓑの本文中における意味として最も適当なものを、次の各群の１～５のうちからそれぞれ一つずつ選び、番号で答えよ。\n\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000１\u3000この上なく集中して\n\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000２\u3000ひたすらまっすぐに\nⓐ\u3000しごく無造作に\u3000\u3000３\u3000きわめてたやすく\n\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000４\u3000すぐに目標を定めて\n\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000５\u3000まったく余裕なしに\n\n\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000１\u3000様々に解釈した\n\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000２\u3000繰り返し考えた\nⓑ\u3000反芻した\u3000\u3000３\u3000好ましく感じた\n\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000４\u3000一つずつ検証した\n\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000５\u3000心に刻んだ\n\n\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000\u3000\n問２\u3000傍線部(ア)とあるが、このときの「私」の心情を九十字以内で説明せよ。\n\u3000\u3000【下書き用】\n\u3000\u3000\n\n\n\n\n問三\u3000傍線部(イ)とあるが、管理人が示した「親しげな態度」とはどのような態度か。九十字以内で説明せよ。\n\u3000\u3000【下書き用】\n\u3000\u3000\n\n\n問四\u3000傍線部(ウ)とあるが、この時の「私」の説明として最も適切なものを、次の１～５のうちから一つ選び、番号で答えよ。\n１\u3000めずらしく親しげな態度を示した管理人との会話が頭を離れず、以前鳥の群れがねぐらに戻るのを見た光景と彼の語った話が一致したことによって、〈翠たち〉の由来に納得がいき、彼らの先祖がいたはずの遠い東南アジアの森に思いを馳せている。\n２\u3000室内に人がいても平気でえさを食べつづけるほど馴れてきた〈翠たち〉が、カメラのレンズには本能的な恐怖心を抱いている様子を見て、野生を受け継いでいることを実感し、植物が鬱蒼と茂ったベランダが彼らの故郷同然であることを願っている。\n２\u3000室内に人がいても平気でえさを食べつづけるほど馴れてきた〈翠たち〉が、カメラのレンズには本能的な恐怖心を抱いている様子を見て、野生を受け継いでいることを実感し、植物が鬱蒼と茂ったベランダが彼らの故郷同然であることを願っている。\n４\u3000以前鳥の群れが飛んでいく光景を見て違和感を受けたが、管理人の話によってそれがインコの群れであったと気付き、熱帯の生息地におけるインコの鳴き声や姿を空想したことで、ベランダに置かれたアロエやハーブが草原のように思われている\n５\u3000管理人との会話が強く印象に残り、ベランダにやってくる〈翠たち〉にも先祖の遺伝子が受け継がれていることに興奮して、熱にうかされた頭の中で、泉の育てた植物が力強く繁茂するベランダに、インコがいる熱帯の森のイメージを重ねている。\n\n\n問五\u3000【文章Ⅰ】の表現に関する説明として適当でないものを、次の１～５のうちから一つ選び、番号で答えよ。\n１\u3000２行目「レタスの葉をもぎとって並べたように」、３行目「ビーズのように」、27・28行目「ゴーガンの風景画から飛びたってきたよう」という直喩を用いることによって、ベランダにやって来るインコの様子を視覚的に印象付けている。\n２\u3000４行目「小首をかしげて」、55行目「天は二物を与えなかった」という慣用表現を用いて、インコの様子や特徴をわかりやすく述べると同時に、鳥を擬人化した語り手や話し手がインコを身近な存在として捉えていることをも示唆している。\n３\u3000６行目「春先から」、26行目「ある土曜日の朝」、34・35行目「管理人に呼びとめられたのはそのころのことである」といったように、本文中では時系列が随所で明確に示され、現在の私の状況に至るまでのいきさつが整理されている。\n４\u300010行目「ギイ、ギイ、ギーッ」という擬音語は、ヤスリをこすり合わせるようなインコのいやな鳴き声を具体化し、62行目「不協和音の一種としか思えないあの鳴き声」とも呼応して、読者に聴覚的イメージを喚起させる効果がある。\n")


    def test_get_style_by_id(self):
        """get_style_by_id のテスト"""
        style_id = "2-10"  # テスト対象の styleId

        # 期待されるスタイルデータ
        expected_style_data = {
            "styleId": "2-10",
            "type": "character",
            "customStyle": "1",
            "name": "2-1_設問_番号 (文字)",
            "basedOn": "DefaultParagraphFont",
            "link": "2-1",
            "font": {
                "ascii": "MS Gothic",
                "hAnsi": "MS Gothic",
                "eastAsia": "MS Gothic",
                "hint": None,
            },
            "color": "ED7D31",
            "size": "21",
        }

        # 関数を実行して結果を取得
        result = get_style_by_id(SAMPLE_DOCX_PATH, style_id)

        # 期待値と比較
        assert result == expected_style_data, (
            f"Expected {expected_style_data}, but got {result}"
        )
if __name__ == '__main__':
    unittest.main()




# 漢数字をアラビア数字に変換する関数のテスト
def test_kanji_number_to_arabic_number():
    test_cases = [
        # 単一の漢数字
        ("〇", "0"),
        ("一", "1"),
        ("二", "2"),
        ("三", "3"),
        ("四", "4"),
        ("五", "5"),
        ("六", "6"),
        ("七", "7"),
        ("八", "8"),
        ("九", "9"),

        # 連続した漢数字
        ("一二三", "123"),
        ("四五六", "456"),
        ("七八九", "789"),
        ("九〇", "90"),

        # 大きな数字
        ("二〇二四", "2024"),  # 西暦表記
        ("五六七八九〇一", "5678901"),

        # 漢数字以外の文字を含むケース
        ("漢字一二三", "123"),  # 文字混じり
        ("テスト五六", "56"),
        ("123四五六", "456"),  # すでに数字が入っている場合

        # 漢数字がない場合
        ("漢字だけ", ""),
        ("", ""),  # 空文字
    ]

    for kanji_input, expected in test_cases:
        assert kanji_number_to_arabic_number(kanji_input) == expected, f"Failed for input: {kanji_input}"
def test_extract_annotation_text_to_list():
    """傍注のリストが正しく抽出されることを確認"""

    def create_mock_paragraph(text):
        """Paragraph モックオブジェクトを作成"""
        mock_p = MagicMock(spec=Paragraph)
        mock_p.text = text
        return mock_p

    # `Paragraph` オブジェクトのリストとしてモックを作成
    mock_document = [
        create_mock_paragraph("これは問題文です。"),
        create_mock_paragraph("これは第二の問題文です。"),
        create_mock_paragraph("（注）１　騒擾――さわぎみだれること。"),  # 傍注の開始
        create_mock_paragraph("　　　２　碁会所や撞球場――「碁会所」は囲碁を打てる場所。「撞球場」はビリヤード場。"),
        create_mock_paragraph("　　　３　元服――昔、貴族や武家の男子が成人することを示した儀式。"),
        create_mock_paragraph("問一　二重傍線部ⓐ・ⓑの本文中における意味として最も適当なものを、次の各群の１～５のうちからそれぞれ一つずつ選び、番号で答えよ。"),  # 収集が停止するべき
    ]

    annotations = extract_annotation_text_to_list(mock_document)

    # デバッグ用に出力
    print(f"Extracted annotations: {annotations}")

    # 期待される結果
    expected = ["騒擾", "碁会所や撞球場", "元服"]

    assert annotations == expected, f"期待される結果: {expected}, 取得した結果: {annotations}"


def test_extract_main_text_and_annotation_to_main_text():
    """問題本文が正しく抽出されることを確認"""

    def create_mock_paragraph(text):
        mock_p = MagicMock(spec=Paragraph)
        mock_p.text = text
        return mock_p

    mock_document = [
        create_mock_paragraph("これは問題文です。"),
        create_mock_paragraph("これは第二の問題文です。"),
        create_mock_paragraph("（注）"),
        create_mock_paragraph("１　漢字――意味の説明"),
        create_mock_paragraph("２　別の単語――別の意味の説明"),
    ]

    main_text = extract_main_text_and_annotation_to_main_text(mock_document)

    # 期待される結果
    expected_texts = ["これは問題文です。", "これは第二の問題文です。"]

    assert len(main_text) == len(expected_texts), "本文の段落数が期待と異なります"

    for i, paragraph in enumerate(main_text):
        assert paragraph.text == expected_texts[i], f"段落 {i} の内容が期待と異なります"
